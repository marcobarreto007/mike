# Copyright (c) 2025-2026 Marco Barreto. All rights reserved.
# Proprietary software - see LICENSE file in project root.

from __future__ import annotations

import re
from pathlib import Path


def _normalize_extracted_text(text: str) -> str:
    lines = []
    for raw_line in str(text or "").splitlines():
        normalized = re.sub(r"\s+", " ", raw_line).strip()
        if normalized:
            lines.append(normalized)
    return "\n".join(lines).strip()


def extract_pdf(path: Path) -> str:
    try:
        import fitz
    except Exception as exc:
        raise RuntimeError("PyMuPDF is required to read PDF knowledge files.") from exc

    pages = []
    with fitz.open(str(path)) as document:
        for page in document:
            page_text = _normalize_extracted_text(page.get_text("text"))
            if page_text:
                pages.append(page_text)
    return "\n\n".join(pages).strip()


def extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is required to read DOCX knowledge files.") from exc

    document = Document(str(path))
    parts = []

    for paragraph in document.paragraphs:
        paragraph_text = _normalize_extracted_text(paragraph.text)
        if paragraph_text:
            parts.append(paragraph_text)

    for table in document.tables:
        for row in table.rows:
            cells = [
                _normalize_extracted_text(cell.text)
                for cell in row.cells
                if _normalize_extracted_text(cell.text)
            ]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_html(path: Path) -> str:
    try:
        from bs4 import BeautifulSoup
    except Exception as exc:
        raise RuntimeError("beautifulsoup4 is required to read HTML knowledge files.") from exc

    raw_html = Path(path).read_text(encoding="utf-8-sig", errors="ignore")
    try:
        soup = BeautifulSoup(raw_html, "lxml")
    except Exception:
        soup = BeautifulSoup(raw_html, "html.parser")

    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    root = soup.find("main") or soup.find("article") or soup.body or soup
    title = ""
    if soup.title and soup.title.string:
        title = _normalize_extracted_text(soup.title.string)

    content = _normalize_extracted_text(root.get_text("\n", strip=True))
    if title and title not in content:
        return f"{title}\n{content}".strip()
    return content
